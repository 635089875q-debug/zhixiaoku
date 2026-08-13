import pytest
from docx import Document
from pypdf import PdfWriter

from app.rag.loader import (
    load_documents,
    load_docx,
    load_pdf,
    load_pdf_pages,
    load_txt,
)

def test_load_txt(tmp_path):
    file_path = tmp_path / "test.txt"

    file_path.write_text(
        "深圳大学位于深圳市",
        encoding="utf-8"
    )

    content = load_txt(file_path)

    assert content == "深圳大学位于深圳市"


def test_load_docx_joins_non_empty_paragraphs(tmp_path):
    file_path = tmp_path / "test.docx"
    document = Document()
    document.add_paragraph("第一段")
    document.add_paragraph("   ")
    document.add_paragraph("第二段")
    document.save(file_path)

    content = load_docx(file_path)

    assert content == "第一段\n\n第二段"


def test_blank_pdf_returns_no_content(tmp_path):
    file_path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)

    with open(file_path, "wb") as file:
        writer.write(file)

    assert load_pdf_pages(file_path) == []
    assert load_pdf(file_path) == ""


def test_encrypted_pdf_is_rejected(tmp_path):
    file_path = tmp_path / "encrypted.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.encrypt("secret")

    with open(file_path, "wb") as file:
        writer.write(file)

    with pytest.raises(
        ValueError,
        match="暂不支持加密PDF文件"
    ):
        load_pdf_pages(file_path)


def test_load_documents_ignores_unsupported_files(tmp_path):
    (tmp_path / "knowledge.txt").write_text(
        "知识库内容",
        encoding="utf-8"
    )
    (tmp_path / "readme.md").write_text(
        "不支持的文件",
        encoding="utf-8"
    )

    documents = load_documents(tmp_path)

    assert documents == [
        {
            "content": "知识库内容",
            "source": "knowledge.txt"
        }
    ]

